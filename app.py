import streamlit as st
import requests
import os
import datetime
from dotenv import load_dotenv
from fpdf import FPDF
from gtts import gTTS
from io import BytesIO
from docx import Document

# Load environment variables
load_dotenv()

API_KEY = os.getenv("API_KEY")
HEADERS = {
    "Authorization": f"Bearer {API_KEY}",
    "HTTP-Referer": "localhost",
    "Content-Type": "application/json"
}
API_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = "mistralai/mixtral-8x7b-instruct"

def generate_openrouter(prompt, max_tokens):
    body = {
        "model": MODEL,
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful and creative story writing assistant. Never repeat or restate the user's first line. Start from the second sentence onward. Always end with a complete sentence."
            },
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7,
        "max_tokens": max_tokens + 30
    }
    try:
        response = requests.post(API_URL, headers=HEADERS, json=body)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        st.error("❌ OpenRouter API call failed")
        st.text(str(e))
        return "[ERROR]"

def convert_to_pdf(title, text, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Times", style="B", size=14)
    pdf.multi_cell(0, 10, title.encode('latin-1', 'replace').decode('latin-1'), align='C')
    pdf.ln(5)
    pdf.set_font("Times", size=12)
    for line in text.split('\n'):
        safe_line = line.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, safe_line)
    pdf.output(filename, "F")

def convert_to_docx(title, story):
    doc = Document()
    doc.add_heading(title, 0)
    for para in story.split("\n\n"):
        doc.add_paragraph(para)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_gtts_audio(text, filename="story.mp3"):
    tts = gTTS(text=text, lang='en')
    tts.save(filename)
    return filename

st.set_page_config(page_title="AI Story Generator", page_icon="📖", layout="wide")

# Light/Dark mode toggle
mode = st.sidebar.radio("🌓 Theme", options=["Dark", "Light"], index=0)
if mode == "Dark":
    st.markdown("""
        <style>
            body, .stApp {
                background-color: #0e1117;
                color: white;
            }
            .stTextInput > label, .stSelectbox > label, .stTextArea > label {
                color: white !important;
            }
            .stButton > button, .stDownloadButton > button {
                background-color: #262730;
                color: white;
                border: 1px solid #555;
                padding: 0.5rem 1rem;
                border-radius: 6px;
            }
            .stButton > button:hover, .stDownloadButton > button:hover {
                background-color: #3c3c3c;
            }
        </style>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
        <style>
            body, .stApp {
                background-color: #ffffff;
                color: black;
            }
            .stTextInput > label, .stSelectbox > label, .stTextArea > label {
                color: black !important;
            }
            .stButton > button, .stDownloadButton > button {
                background-color: #f2f2f2;
                color: black;
                border: 1px solid #ccc;
                padding: 0.5rem 1rem;
                border-radius: 6px;
            }
            .stButton > button:hover, .stDownloadButton > button:hover {
                background-color: #e6e6e6;
            }
        </style>
    """, unsafe_allow_html=True)

st.title("📖 AI Story Generator")
st.markdown("Craft immersive stories with setting and tone.")

if "story" not in st.session_state:
    st.session_state.story = ""
if "title" not in st.session_state:
    st.session_state.title = ""
if "history" not in st.session_state:
    st.session_state.history = []
if "continue_clicked" not in st.session_state:
    st.session_state.continue_clicked = False
if "reset_flag" not in st.session_state:
    st.session_state.reset_flag = False

st.sidebar.title("📚 Story History")

selected_history = None
if st.session_state.history:
    st.sidebar.markdown("""
    <p style='font-weight:600; font-size:16px; color:white; margin-top: 10px;'>📜 View previous story</p>
    """, unsafe_allow_html=True)

    titles = [f"{len(st.session_state.history)-i}. {h['title']}" for i, h in enumerate(st.session_state.history[::-1])]
    selected = st.sidebar.selectbox("", ["Select one"] + titles)
    if selected != "Select one":
        idx = int(selected.split(". ")[0]) - 1
        selected_history = st.session_state.history[idx]
        st.session_state.story = selected_history["story"]
        st.session_state.title = selected_history["title"]
        st.session_state.continue_clicked = False
else:
    st.sidebar.write("No stories yet.")


# Inputs
genre = st.selectbox("🎭 Choose a genre", ["Select a genre","Fantasy", "Sci-Fi", "Romance", "Mystery", "Comedy", "Horror"], index=0, key="genre")
character = st.text_input("🧑 Character name (optional)", placeholder="e.g. Elena", key="character")
setting_choice = st.selectbox("🌍 Choose a setting", ["Select a setting","Bookstore","Town","Village","Party","Forest", "Café", "Castle", "Space Station", "Island", "Unknown"], index=0, key="setting_choice")
mood = st.selectbox("🎨 Select story tone", ["Select a story tone","Light", "Whimsical", "Mysterious", "Dramatic", "Dark"], index=0, key="mood")
emotion = st.selectbox("💓 Character Emotion", ["Select your character emotion","Neutral", "Happy", "Nervous", "Confident", "Sad", "Angry", "Excited", "Scared"], index=0, key="emotion")
start_prompt = st.text_area("✍ Write the first line of your story", height=100, placeholder="e.g. Elena was standing in a corner, nervously fidgeting with her hands...", key="start_prompt")

length = st.selectbox("🧱 Story Length", ["Short", "Medium", "Long"], index=1, key="length")
length_map = {"Short": 300, "Medium": 600, "Long": 800}
max_tokens = length_map[length]

# --- Initial Story Generation ---
if start_prompt and st.button("🚀 Generate Initial Story"):
    with st.spinner("✨ Generating immersive setting and story..."):
        setting_prompt = f"Create a vivid 2-3 line world-building description of a {setting_choice.lower()} in a {genre} story with a {mood.lower()} tone."
        setting_description = generate_openrouter(setting_prompt, 100)

        story_prompt = f"""
Continue this {genre} story with a {mood.lower()} tone. Use this first line exactly once and do not repeat it. Begin directly from the second sentence onward.

First line: {start_prompt}
Setting: {setting_description}
Character emotion: {emotion}
"""
        story_start = generate_openrouter(story_prompt, max_tokens)
        st.session_state.story = setting_description + "\n\n" + story_start.strip()

        title_prompt = f"Suggest only a short and creative title for the following {genre} story. Return just the title. Do not explain or describe it.\n\n{st.session_state.story}"
        st.session_state.title = generate_openrouter(title_prompt, 30)

if st.session_state.story:
    if st.session_state.title:
        st.subheader(f"📖 {st.session_state.title}")
    st.markdown("### 📚 Your Story So Far")
    for para in st.session_state.story.split("\n\n"):
        st.markdown(para)

    if st.button("📖 Continue Story"):
        st.session_state.continue_clicked = True

    if st.session_state.continue_clicked:
        with st.spinner("✍ Continuing story..."):
            continuation_prompt = f"Continue the story below with a new paragraph in the {genre} genre with a {mood.lower()} tone. The main character feels {emotion.lower()}. Do not repeat any lines. End with a complete sentence.\n\n{st.session_state.story}"
            next_paragraph = generate_openrouter(continuation_prompt, max_tokens)
            st.session_state.story += "\n\n" + next_paragraph.strip()
            st.session_state.continue_clicked = False
        st.rerun()

    if st.button("🔊 Speak the Story"):
        with st.spinner("🔈 Generating voiceover..."):
            audio_file = generate_gtts_audio(st.session_state.story)
            with open(audio_file, "rb") as audio:
                st.audio(audio.read(), format="audio/mp3")

    docx_file = convert_to_docx(st.session_state.title, st.session_state.story)
    st.download_button(
        label="📥 Download Story as DOC",
        data=docx_file,
        file_name=f"{st.session_state.title}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    pdf_filename = "your_story.pdf"
    convert_to_pdf(st.session_state.title, st.session_state.story, pdf_filename)
    with open(pdf_filename, "rb") as pdf_file:
        st.download_button(
            label="📄 Download Story as PDF",
            data=pdf_file,
            file_name=pdf_filename,
            mime="application/pdf"
        )

if st.button("🔄 Start Over"):
    if st.session_state.title and st.session_state.story and all(h["title"] != st.session_state.title for h in st.session_state.history):
        st.session_state.history.append({
            "title": st.session_state.title,
            "story": st.session_state.story,
            "genre": genre,
            "time": str(datetime.datetime.now())
        })
    st.session_state.reset_flag = True

if st.session_state.reset_flag:
    st.session_state["story"] = ""
    st.session_state["title"] = ""
    st.session_state["continue_clicked"] = False
    st.session_state.reset_flag = False

    st.experimental_set_query_params()
    st.rerun()